import os
from datetime import date
from typing import List

from flask import (
    Flask, render_template, request, redirect, url_for,
    jsonify, flash, send_file, abort
)
from flask_sqlalchemy import SQLAlchemy

# ======== DOCX helpers ========
try:
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    _DOCX_ENABLED = True
except Exception:
    _DOCX_ENABLED = False

# ---------------- Config ----------------
ITEMS_TOTAL = 40
DISCONTINUE_THRESHOLD = 28   # >= 28 => discontinue
ONE_BELOW_THRESHOLD = 16     # 16..27 => start one below; 0..15 => two below

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret")
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///philiri.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# ---------------- Models ----------------
class Class(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    teacher = db.Column(db.String(120), default="")
    school = db.Column(db.String(120), default="")
    grade = db.Column(db.Integer, default=7)
    section = db.Column(db.String(80), default="")
    screening_level_eng = db.Column(db.String(20), default="GST")
    screening_level_fil = db.Column(db.String(20), default="GST")
    date_text = db.Column(db.String(40), default="")

class Learner(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    class_id = db.Column(db.Integer, db.ForeignKey("class.id"), nullable=False)
    name = db.Column(db.String(120), default="")
    gender = db.Column(db.String(1), default="M")
    took_eng = db.Column(db.Boolean, default=True)
    took_fil = db.Column(db.Boolean, default=True)

    eng_literal = db.Column(db.Integer, default=0)
    eng_inferential = db.Column(db.Integer, default=0)
    eng_critical = db.Column(db.Integer, default=0)

    fil_literal = db.Column(db.Integer, default=0)
    fil_inferential = db.Column(db.Integer, default=0)
    fil_critical = db.Column(db.Integer, default=0)

# -------------- Logic helpers --------------
def clamp(n, lo, hi):
    return max(lo, min(hi, n))

def compute_total(lit, inf, cri):
    return int(lit or 0) + int(inf or 0) + int(cri or 0)

def starting_point_for(grade: int, total: int) -> str:
    if total >= DISCONTINUE_THRESHOLD:
        return "DISCONTINUE"
    base = max(1, int(grade) - 1)
    start_level = max(1, base - 1 if total >= ONE_BELOW_THRESHOLD else base - 2)
    return f"Grade {start_level}"

def _gender_bucket(g: str) -> int:
    g = (g or "").strip().upper()
    if g.startswith("M"):
        return 0
    if g.startswith("F"):
        return 1
    return 2

def _sort_rows(rows: List[dict]) -> List[dict]:
    return sorted(
        rows,
        key=lambda r: (_gender_bucket(r.get("gender", "")),
                       (r.get("name", "") or "").strip().upper())
    )

# ======== Template paths ========
def _template_path_eng() -> str:
    here = os.path.dirname(os.path.abspath(__file__))
    a = os.path.join(here, "static", "docs", "Portrait_Header-Footer.docx")
    b = "/mnt/data/Portrait_Header-Footer.docx"
    return a if os.path.isfile(a) else b

def _template_path_fil() -> str:
    here = os.path.dirname(os.path.abspath(__file__))
    a = os.path.join(here, "static", "docs", "Portrait_Header-FIL.docx")
    b = "/mnt/data/Portrait_Header-FIL.docx"
    return a if os.path.isfile(a) else b

# ======== Tight paragraph helpers ========
def _set_p_no_space(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1

def _set_cell_paras_no_space(cell):
    for p in cell.paragraphs:
        _set_p_no_space(p)

def _docx_set_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'Book Antiqua'
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
    except Exception:
        pass
    style.font.size = Pt(10)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1

# ======== Helpers that write into template tables ========
def _cell_set_text(cell, text, *, bold=False, underline=False, align=None, upper=True):
    """
    Write text into a table cell using one tight paragraph and one run.
    - bold: for labels only
    - underline: for values only (text underline, not cell border)
    - align: optional paragraph alignment
    - upper: values often uppercase to match samples
    """
    # Don't wipe existing styled labels unless we're writing a value.
    cell.text = ""
    p = cell.paragraphs[0]
    _set_p_no_space(p)
    if align is not None:
        p.alignment = align
    run_text = (text or "")
    run_text = run_text.upper() if upper else run_text
    r = p.add_run(run_text)
    r.bold = bool(bold)
    if underline:
        r.underline = True

def _norm(s: str) -> str:
    return (s or "").strip().lower()

def _looks_like_meta_table(tbl) -> bool:
    try:
        if len(tbl.rows) >= 3 and len(tbl.columns) >= 4:
            a = _norm(tbl.cell(0, 0).text)   # left label row 1
            b = _norm(tbl.cell(0, 2).text)   # right label row 1
            # English OR Filipino labels
            return (
                (a.startswith("teacher") and b.startswith("grade")) or
                (a.startswith("guro") and b.startswith("antas"))
            )
    except Exception:
        pass
    return False

def _looks_like_results_table(tbl) -> bool:
    try:
        hdr = " ".join((c.text or "").strip().upper() for c in tbl.rows[0].cells)
        has_name = ("NAME" in hdr) or ("PANGALAN" in hdr)
        has_gender = ("GENDER" in hdr) or ("KASARIAN" in hdr)
        has_score = ("SCORE" in hdr) or ("MARKA" in hdr)
        return has_name and has_gender and has_score
    except Exception:
        return False

def _fill_template_tables(
    doc: Document,
    cls_dict: dict,
    rows: List[dict],
    *,
    type_of_test_text: str
) -> bool:
    """
    Populate the meta table and the results table that already exist in the template.
    We DO NOT touch label cells, to preserve any color/formatting in the template
    (especially the Filipino red labels). We only write the VALUE cells:
      row0 col1/3, row1 col1/3, row2 col1/3
    """
    meta_tbl = None
    results_tbl = None
    for t in doc.tables:
        if meta_tbl is None and _looks_like_meta_table(t):
            meta_tbl = t
        elif results_tbl is None and _looks_like_results_table(t):
            results_tbl = t
        if meta_tbl and results_tbl:
            break

    if meta_tbl is None or results_tbl is None:
        return False

    # --- Meta values (underlined, not bold) ---
    _cell_set_text(meta_tbl.cell(0, 1), cls_dict.get("teacher", ""), underline=True)
    _cell_set_text(meta_tbl.cell(0, 3), str(cls_dict.get("grade", "")), underline=True)

    _cell_set_text(meta_tbl.cell(1, 1), cls_dict.get("school", ""), underline=True)
    _cell_set_text(meta_tbl.cell(1, 3), cls_dict.get("section", ""), underline=True)

    _cell_set_text(meta_tbl.cell(2, 1), type_of_test_text, underline=True)
    _cell_set_text(meta_tbl.cell(2, 3), cls_dict.get("date_text", ""), underline=True)

    for r in meta_tbl.rows:
        for c in r.cells:
            _set_cell_paras_no_space(c)

    # --- Results table (reuse header already in template) ---
    start_row_index = 1  # row 0 is header
    needed = start_row_index + max(1, len(rows))
    while len(results_tbl.rows) < needed:
        results_tbl.add_row()

    for i, rec in enumerate(rows, start=1):
        cells = results_tbl.rows[start_row_index + (i - 1)].cells
        _cell_set_text(cells[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER, upper=False)
        _cell_set_text(cells[1], rec.get("name", ""))
        _cell_set_text(cells[2], rec.get("gender", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_set_text(cells[3], str(rec.get("score", "")), align=WD_ALIGN_PARAGRAPH.CENTER, upper=False)
        _cell_set_text(cells[4], rec.get("start", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        for c in cells:
            _set_cell_paras_no_space(c)

    # Blank out extra template rows (keep borders)
    for r in results_tbl.rows[start_row_index + len(rows):]:
        for c in r.cells:
            _cell_set_text(c, "", upper=False)

    for c in results_tbl.rows[0].cells:
        _set_cell_paras_no_space(c)

    return True

# ======== Fallback builder (only if template tables are missing) ========
def _fallback_build(doc: Document, cls_dict: dict, rows: List[dict], *, title1: str, title2: str, type_of_test_text: str):
    # Titles (tight)
    p1 = doc.add_paragraph(); _set_p_no_space(p1)
    r1 = p1.add_run(title1); r1.bold = True; r1.font.size = Pt(14)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(); _set_p_no_space(p2)
    r2 = p2.add_run(title2); r2.bold = True; r2.font.size = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta (labels bold; values underlined only)
    t = doc.add_table(rows=3, cols=4); t.autofit = True
    labels = [("Teacher:", "Grade:"), ("School:", "Section:"), ("Type of Test:", "Date:")]
    vals = [
        (cls_dict.get("teacher", ""), str(cls_dict.get("grade", ""))),
        (cls_dict.get("school", ""), cls_dict.get("section", "")),
        (type_of_test_text, cls_dict.get("date_text", "")),
    ]
    for r in range(3):
        _cell_set_text(t.cell(r, 0), labels[r][0], bold=True, upper=False)
        _cell_set_text(t.cell(r, 1), vals[r][0], underline=True)
        _cell_set_text(t.cell(r, 2), labels[r][1], bold=True, upper=False)
        _cell_set_text(t.cell(r, 3), vals[r][1], underline=True)

    # Results
    rt = doc.add_table(rows=1, cols=5)
    headers = ("#", "NAME", "GENDER", "SCORE", "START LEVEL OF GRADE PASSAGE")
    for i, h in enumerate(headers):
        _cell_set_text(rt.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, upper=False)
    for i, rec in enumerate(rows, start=1):
        row = rt.add_row().cells
        _cell_set_text(row[0], str(i), align=WD_ALIGN_PARAGRAPH.CENTER, upper=False)
        _cell_set_text(row[1], rec.get("name", ""))
        _cell_set_text(row[2], rec.get("gender", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_set_text(row[3], str(rec.get("score", "")), align=WD_ALIGN_PARAGRAPH.CENTER, upper=False)
        _cell_set_text(row[4], rec.get("start", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

# ======== Builders for ENG / FIL ========
def _build_gst_docx_from_template(
    cls_dict: dict,
    rows: List[dict],
    *,
    template_path: str,
    title1: str,
    title2: str,
    type_of_test_text: str
) -> BytesIO:
    if not _DOCX_ENABLED:
        abort(500, description="python-docx is not installed. Run: pip install python-docx")
    if not os.path.isfile(template_path):
        abort(500, description=f"Word template not found: {template_path}")

    rows = _sort_rows(rows)  # always sort before writing

    doc = Document(template_path)  # keeps header/footer/art & existing tables
    _docx_set_defaults(doc)

    if not _fill_template_tables(doc, cls_dict, rows, type_of_test_text=type_of_test_text):
        _fallback_build(doc, cls_dict, rows, title1=title1, title2=title2, type_of_test_text=type_of_test_text)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def _build_gst_docx_eng(cls_dict: dict, rows: List[dict]) -> BytesIO:
    return _build_gst_docx_from_template(
        cls_dict, rows,
        template_path=_template_path_eng(),
        title1="ENGLISH GST RESULTS",
        title2="STAGE 2 ADMISSION IN PHIL-IRI",
        type_of_test_text="Screening Test Level (English)"
    )

def _build_gst_docx_fil(cls_dict: dict, rows: List[dict]) -> BytesIO:
    return _build_gst_docx_from_template(
        cls_dict, rows,
        template_path=_template_path_fil(),
        title1="FILIPINO GST RESULTS",
        title2="STAGE 2 ADMISSION IN PHIL-IRI",
        type_of_test_text="Screening Test Level (Filipino)"
    )

# ---------------- Routes ----------------
def init_db():
    db.create_all()
    if Class.query.count() == 0:
        c = Class(
            teacher="", school="", grade=7, section="",
            screening_level_eng="GST", screening_level_fil="GST",
            date_text=str(date.today()),
        )
        db.session.add(c)
        db.session.commit()

@app.get("/")
def home():
    classes = Class.query.order_by(Class.id.desc()).all()
    return render_template("home.html", classes=classes)

@app.post("/class/new")
def class_new():
    c = Class(
        teacher=request.form.get("teacher", "").strip(),
        school=request.form.get("school", "").strip(),
        grade=clamp(int(request.form.get("grade", 7) or 7), 1, 12),
        section=request.form.get("section", "").strip(),
        screening_level_eng=(request.form.get("screening_level_eng") or "").strip(),
        screening_level_fil=(request.form.get("screening_level_fil") or "").strip(),
        date_text=request.form.get("date_text", "").strip(),
    )
    db.session.add(c)
    db.session.commit()
    return redirect(url_for("class_edit", class_id=c.id))

@app.get("/class/<int:class_id>")
def class_edit(class_id):
    c = Class.query.get_or_404(class_id)
    learners = (
        Learner.query.filter_by(class_id=class_id)
        .order_by(Learner.gender.desc(), Learner.id.asc())
        .all()
    )
    return render_template(
        "class_edit.html",
        cls=c,
        learners=learners,
        items_total=ITEMS_TOTAL,
        disc=DISCONTINUE_THRESHOLD,
        one_below=ONE_BELOW_THRESHOLD,
    )

@app.post("/class/<int:class_id>/meta")
def class_update_meta(class_id):
    c = Class.query.get_or_404(class_id)
    c.teacher = request.form.get("teacher", "").strip()
    c.school = request.form.get("school", "").strip()
    c.grade = clamp(int(request.form.get("grade", c.grade or 7) or 7), 1, 12)
    c.section = request.form.get("section", "").strip()
    c.screening_level_eng = (request.form.get("screening_level_eng") or "").strip()
    c.screening_level_fil = (request.form.get("screening_level_fil") or "").strip()
    c.date_text = request.form.get("date_text", "").strip()
    db.session.commit()
    flash("Saved class details.", "ok")
    return redirect(url_for("class_edit", class_id=class_id))

@app.post("/api/class/<int:class_id>/learners/save")
def api_learners_save(class_id):
    _ = Class.query.get_or_404(class_id)
    rows = request.get_json(force=True).get("rows", [])
    Learner.query.filter_by(class_id=class_id).delete()
    for r in rows:
        L = Learner(
            class_id=class_id,
            name=(r.get("name", "") or "").strip(),
            gender=(r.get("gender", "M") or "M")[:1].upper(),
            took_eng=bool(r.get("took_eng", True)),
            took_fil=bool(r.get("took_fil", True)),
            eng_literal=int(r.get("eng_literal", 0) or 0),
            eng_inferential=int(r.get("eng_inferential", 0) or 0),
            eng_critical=int(r.get("eng_critical", 0) or 0),
            fil_literal=int(r.get("fil_literal", 0) or 0),
            fil_inferential=int(r.get("fil_inferential", 0) or 0),
            fil_critical=int(r.get("fil_critical", 0) or 0),
        )
        db.session.add(L)
    db.session.commit()
    return jsonify({"ok": True})

def split_by_gender(learners: List[Learner]):
    males = [x for x in learners if (x.gender or "M").upper().startswith("M")]
    females = [x for x in learners if (x.gender or "F").upper().startswith("F")]
    return males, females

@app.get("/form1a/<int:class_id>")
def form1a(class_id):
    c = Class.query.get_or_404(class_id)
    learners = (
        Learner.query.filter_by(class_id=class_id)
        .order_by(Learner.gender.desc(), Learner.id.asc())
        .all()
    )
    m, f = split_by_gender(learners)
    return render_template(
        "form1a.html",
        cls=c,
        males=m,
        females=f,
        items_total=ITEMS_TOTAL,
        disc=DISCONTINUE_THRESHOLD,
        one_below=ONE_BELOW_THRESHOLD,
        compute_total=compute_total,
        starting_point_for=starting_point_for,
    )

@app.get("/form1b/<int:class_id>")
def form1b(class_id):
    c = Class.query.get_or_404(class_id)
    learners = (
        Learner.query.filter_by(class_id=class_id)
        .order_by(Learner.gender.desc(), Learner.id.asc())
        .all()
    )
    m, f = split_by_gender(learners)
    return render_template(
        "form1b.html",
        cls=c,
        males=m,
        females=f,
        items_total=ITEMS_TOTAL,
        disc=DISCONTINUE_THRESHOLD,
        one_below=ONE_BELOW_THRESHOLD,
        compute_total=compute_total,
        starting_point_for=starting_point_for,
    )

# --------- GST ENGLISH ----------
@app.get("/gst/en/<int:class_id>")
def gst_en(class_id):
    c = Class.query.get_or_404(class_id)
    learners = Learner.query.filter_by(class_id=class_id).order_by(Learner.name.asc()).all()
    rows = []
    for s in learners:
        if not s.took_eng:
            continue
        total = compute_total(s.eng_literal, s.eng_inferential, s.eng_critical)
        sp = starting_point_for(c.grade, total)
        if sp != "DISCONTINUE":
            rows.append({"name": s.name, "gender": s.gender, "score": total, "start": sp})
    rows = _sort_rows(rows)
    return render_template("gst_en.html", cls=c, rows=rows)

@app.get("/gst/en/<int:class_id>/export")
def export_gst_en_docx(class_id):
    c = Class.query.get_or_404(class_id)
    learners = Learner.query.filter_by(class_id=class_id).order_by(Learner.name.asc()).all()
    rows = []
    for s in learners:
        if not s.took_eng:
            continue
        total = compute_total(s.eng_literal, s.eng_inferential, s.eng_critical)
        sp = starting_point_for(c.grade, total)
        if sp != "DISCONTINUE":
            rows.append({"name": s.name, "gender": s.gender, "score": total, "start": sp})
    rows = _sort_rows(rows)
    cls_dict = {
        "teacher": c.teacher, "school": c.school, "grade": c.grade,
        "section": c.section, "date_text": c.date_text
    }
    bio = _build_gst_docx_eng(cls_dict, rows)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"ENGLISH_GST_RESULTS_{class_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# --------- GST FILIPINO ----------
@app.get("/gst/fil/<int:class_id>")
def gst_fil(class_id):
    c = Class.query.get_or_404(class_id)
    learners = Learner.query.filter_by(class_id=class_id).order_by(Learner.name.asc()).all()
    rows = []
    for s in learners:
        if not s.took_fil:
            continue
        total = compute_total(s.fil_literal, s.fil_inferential, s.fil_critical)
        sp = starting_point_for(c.grade, total)
        if sp != "DISCONTINUE":
            rows.append({"name": s.name, "gender": s.gender, "score": total, "start": sp})
    rows = _sort_rows(rows)
    return render_template("gst_fil.html", cls=c, rows=rows)

@app.get("/gst/fil/<int:class_id>/export")
def export_gst_fil_docx(class_id):
    c = Class.query.get_or_404(class_id)
    learners = Learner.query.filter_by(class_id=class_id).order_by(Learner.name.asc()).all()
    rows = []
    for s in learners:
        if not s.took_fil:
            continue
        total = compute_total(s.fil_literal, s.fil_inferential, s.fil_critical)
        sp = starting_point_for(c.grade, total)
        if sp != "DISCONTINUE":
            rows.append({"name": s.name, "gender": s.gender, "score": total, "start": sp})
    rows = _sort_rows(rows)
    cls_dict = {
        "teacher": c.teacher, "school": c.school, "grade": c.grade,
        "section": c.section, "date_text": c.date_text
    }
    bio = _build_gst_docx_fil(cls_dict, rows)
    return send_file(
        bio,
        as_attachment=True,
        download_name=f"FILIPINO_GST_RESULTS_{class_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ======== Legacy POST (kept for compatibility; uses English template) ========
@app.post("/export-gst-docx")
def export_gst_docx_post():
    data = request.get_json(force=True) or {}
    cls_dict = data.get("cls", {}) or {}
    rows = data.get("rows", []) or []
    rows = _sort_rows(rows)
    bio = _build_gst_docx_eng(cls_dict, rows)
    return send_file(
        bio,
        as_attachment=True,
        download_name="ENGLISH_GST_RESULTS.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ---------------------------------------
if __name__ == "__main__":
    with app.app_context():
        init_db()
    app.run(debug=True, host="0.0.0.0", port=5001)
